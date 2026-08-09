#!/usr/bin/env python3
"""Deterministic converter: QMatch IQ DOCX v2 (340) -> iq_bank_tr_v1.json.

Primary source: docs/source/assessment/iq/QMatch_Bilissel_Muhakeme_Soru_Bankasi_v2_340.docx
Does not generate or rewrite question content.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx required. Use .venv_iq_convert/bin/python after "
        "`python3.12 -m venv .venv_iq_convert && pip install python-docx pypdf`"
    ) from exc

PARSER_VERSION = "convert_iq_bank_v2.py@1.0.1"

SECTION_MAP = {
    "Bölüm 1 - Mantıksal Muhakeme": "logical_reasoning",
    "Bölüm 2 - Örüntü Muhakemesi": "pattern_reasoning",
    "Bölüm 3 - Sözel Muhakeme": "verbal_reasoning",
    "Bölüm 4 - Uzamsal Muhakeme": "spatial_reasoning",
}

QUESTION_RE = re.compile(r"^(\d+)\.\s+(.*)$", re.DOTALL)
# Optional trailing marker uses source wording "v2 yeniden yazım" (not past tense).
ID_RE = re.compile(
    r"^ID:\s*(\S+)\s*\|\s*Aile:\s*(\S+)(?:\s*\|\s*v2 yeniden yazım)?\s*$",
    re.IGNORECASE,
)
OPTION_RE = re.compile(r"^([A-D])\)\s*(.*)$", re.DOTALL)
REWRITTEN_LINE_RE = re.compile(r"^([a-z0-9_]+)\s+-\s+(iq_template_\d+)\s*$")
CONTROL_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")
INLINE_REWRITTEN_RE = re.compile(r"\|\s*v2 yeniden yazım\s*$", re.IGNORECASE)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def nonempty_paragraphs(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]


def parse_answer_tables(doc: Document) -> dict[str, str]:
    """Map item_id -> uppercase answer letter from answer-key tables."""
    answers: dict[str, str] = {}
    # Tables 1..4 are dimension answer keys (after summary table 0).
    if len(doc.tables) < 5:
        raise SystemExit(f"Expected >=5 tables, found {len(doc.tables)}")
    for table in doc.tables[1:5]:
        rows = table.rows
        header = [c.text.strip() for c in rows[0].cells]
        if header[:4] != ["No", "Soru ID", "Aile", "Yanıt"]:
            raise SystemExit(f"Unexpected answer table header: {header}")
        for row in rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 4:
                raise SystemExit(f"Malformed answer row: {cells}")
            item_id, letter = cells[1], cells[3].strip().upper()
            if letter not in {"A", "B", "C", "D"}:
                raise SystemExit(f"Invalid answer letter for {item_id}: {letter}")
            if item_id in answers:
                raise SystemExit(f"Duplicate answer key for {item_id}")
            answers[item_id] = letter
    if len(answers) != 340:
        raise SystemExit(f"Expected 340 answer keys, found {len(answers)}")
    return answers


def parse_rewritten_ids(paras: list[str]) -> set[str]:
    try:
        start = paras.index("v2 Yeniden Yazılan Sorular")
    except ValueError as exc:
        raise SystemExit("Missing rewritten section header") from exc
    end = paras.index("Otomatik Bütünlük Kontrolü")
    ids: set[str] = set()
    for line in paras[start + 1 : end]:
        m = REWRITTEN_LINE_RE.match(line)
        if m:
            ids.add(m.group(1))
    if len(ids) != 40:
        raise SystemExit(f"Expected 40 rewritten ids, found {len(ids)}: {sorted(ids)}")
    return ids


def parse_question_blocks(
    paras: list[str],
) -> list[dict]:
    """Extract question blocks with dimension from section headers."""
    try:
        start = next(i for i, p in enumerate(paras) if p.startswith("Bölüm 1 -"))
        end = paras.index("Cevap Anahtarı")
    except (StopIteration, ValueError) as exc:
        raise SystemExit("Could not locate question body bounds") from exc

    dimension = None
    blocks: list[dict] = []
    i = start
    body = paras[start:end]
    idx = 0
    while idx < len(body):
        line = body[idx]
        if line in SECTION_MAP:
            dimension = SECTION_MAP[line]
            idx += 1
            # skip optional section blurb
            if idx < len(body) and not QUESTION_RE.match(body[idx]) and not body[idx].startswith("Bölüm "):
                idx += 1
            continue
        m = QUESTION_RE.match(line)
        if not m:
            raise SystemExit(
                f"Unexpected paragraph in question body at local index {idx}: {line[:120]!r}"
            )
        if dimension is None:
            raise SystemExit("Question encountered before dimension section")
        source_order = int(m.group(1))
        prompt = m.group(2).strip()
        idx += 1
        if idx >= len(body):
            raise SystemExit(f"Incomplete block after question {source_order}")
        id_m = ID_RE.match(body[idx])
        if not id_m:
            raise SystemExit(
                f"Missing ID line after question {source_order}: {body[idx][:120]!r}"
            )
        item_id, family = id_m.group(1), id_m.group(2)
        marked_inline = bool(INLINE_REWRITTEN_RE.search(body[idx]))
        idx += 1
        options: dict[str, str] = {}
        for expected in ("A", "B", "C", "D"):
            if idx >= len(body):
                raise SystemExit(f"Missing option {expected} for {item_id}")
            om = OPTION_RE.match(body[idx])
            if not om or om.group(1) != expected:
                raise SystemExit(
                    f"Expected option {expected} for {item_id}, got: {body[idx][:120]!r}"
                )
            options[expected] = om.group(2).strip()
            idx += 1
        if not prompt:
            raise SystemExit(f"Empty prompt for {item_id}")
        if CONTROL_RE.search(prompt) or any(CONTROL_RE.search(v) for v in options.values()):
            raise SystemExit(f"Control characters in {item_id}")
        blocks.append(
            {
                "source_order": source_order,
                "id": item_id,
                "template_family_id": family,
                "dimension": dimension,
                "prompt": prompt,
                "options_upper": options,
                "marked_inline_rewritten": marked_inline,
            }
        )
    if len(blocks) != 340:
        raise SystemExit(f"Expected 340 question blocks, found {len(blocks)}")
    return blocks


def build_bank(
    *,
    docx_path: Path,
    source_name: str,
) -> dict:
    doc = Document(str(docx_path))
    paras = nonempty_paragraphs(doc)
    answers = parse_answer_tables(doc)
    rewritten = parse_rewritten_ids(paras)
    blocks = parse_question_blocks(paras)

    # Associate answers by ID
    missing_answers = [b["id"] for b in blocks if b["id"] not in answers]
    if missing_answers:
        raise SystemExit(f"Missing answers for: {missing_answers[:10]}")
    extra = set(answers) - {b["id"] for b in blocks}
    if extra:
        raise SystemExit(f"Answer keys without blocks: {sorted(extra)[:10]}")

    # Prefer explicit rewritten list; cross-check inline markers for integrity.
    inline_rewritten = {b["id"] for b in blocks if b["marked_inline_rewritten"]}
    if inline_rewritten != rewritten:
        only_list = sorted(rewritten - inline_rewritten)
        only_inline = sorted(inline_rewritten - rewritten)
        raise SystemExit(
            "Rewritten list vs inline markers mismatch: "
            f"only_list={only_list[:10]} only_inline={only_inline[:10]}"
        )

    items = []
    for b in sorted(blocks, key=lambda x: x["source_order"]):
        letter = answers[b["id"]]
        options = [
            {"id": k.lower(), "text": b["options_upper"][k]}
            for k in ("A", "B", "C", "D")
        ]
        items.append(
            {
                "id": b["id"],
                "dimension": b["dimension"],
                "template_family_id": b["template_family_id"],
                "prompt": b["prompt"],
                "options": options,
                "correct_option_id": letter.lower(),
                "source_order": b["source_order"],
                "revision_status": (
                    "rewritten_v2" if b["id"] in rewritten else "retained_v2"
                ),
                "review_status": "desk_reviewed_candidate",
            }
        )

    return {
        "schema_version": "qmatch_iq_bank_v1",
        "bank_version": "tr_v2_340",
        "locale": "tr-TR",
        "source": source_name,
        "parser_version": PARSER_VERSION,
        "items": items,
    }


def dumps_deterministic(bank: dict) -> str:
    return json.dumps(bank, ensure_ascii=False, indent=2) + "\n"


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--docx",
        default="docs/source/assessment/iq/QMatch_Bilissel_Muhakeme_Soru_Bankasi_v2_340.docx",
    )
    parser.add_argument(
        "--out",
        default="assets/data/assessment_v3/iq/iq_bank_tr_v1.json",
    )
    parser.add_argument(
        "--meta-out",
        default="assets/data/assessment_v3/iq/reports/iq_bank_tr_v1_conversion_meta.json",
    )
    args = parser.parse_args()

    root = Path.cwd()
    docx_path = root / args.docx
    out_path = root / args.out
    meta_path = root / args.meta_out
    if not docx_path.exists():
        print(f"MISSING_SOURCE {docx_path}", file=sys.stderr)
        return 2

    bank = build_bank(
        docx_path=docx_path,
        source_name="QMatch_Bilissel_Muhakeme_Soru_Bankasi_v2_340",
    )
    text = dumps_deterministic(bank)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(text, encoding="utf-8")

    dims = Counter(i["dimension"] for i in bank["items"])
    families = Counter(i["template_family_id"] for i in bank["items"])
    rewritten = Counter(
        i["dimension"]
        for i in bank["items"]
        if i["revision_status"] == "rewritten_v2"
    )
    answers = Counter(i["correct_option_id"] for i in bank["items"])
    meta = {
        "parser_version": PARSER_VERSION,
        "docx_path": str(docx_path.as_posix()),
        "docx_sha256": sha256_file(docx_path),
        "output_path": str(out_path.as_posix()),
        "output_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
        "item_count": len(bank["items"]),
        "dimension_counts": dict(dims),
        "family_count": len(families),
        "rewritten_counts": dict(rewritten),
        "answer_position_counts": dict(answers),
        "content_generated": False,
        "content_rewritten_by_converter": False,
    }
    meta_path.parent.mkdir(parents=True, exist_ok=True)
    meta_path.write_text(
        json.dumps(meta, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    print(json.dumps({"ok": True, **meta}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
